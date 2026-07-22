#!/usr/bin/env python3
"""Export Speciedex records to Microsoft Word DOCX."""

from __future__ import annotations

import argparse
import json
from collections.abc import Sequence
from pathlib import Path

from export_common import (
    add_common_arguments,
    connect_readonly,
    ensure_parent,
    iter_rows,
    preferred_fields,
    validate_common_arguments,
)

DEFAULT_OUTPUT = Path("static/data/exports/speciedex.docx")


def parse_args(argv: Sequence[str] | None = None) -> argparse.Namespace:
    parser = argparse.ArgumentParser(
        description="Export Speciedex records to DOCX."
    )
    add_common_arguments(parser, default_output=DEFAULT_OUTPUT)
    parser.add_argument("--page-break-every", type=int, default=0)
    return parser.parse_args(argv)


def main(argv: Sequence[str] | None = None) -> int:
    args = parse_args(argv)
    validate_common_arguments(args)
    ensure_parent(args.output)

    try:
        from docx import Document
        from docx.shared import Inches
    except ImportError as error:
        raise SystemExit(
            "DOCX export requires python-docx: python -m pip install python-docx"
        ) from error

    document = Document()
    document.core_properties.title = args.title
    document.core_properties.author = "Speciedex.org"
    document.add_heading(args.title, level=0)

    section = document.sections[0]
    section.top_margin = Inches(0.55)
    section.bottom_margin = Inches(0.55)
    section.left_margin = Inches(0.55)
    section.right_margin = Inches(0.55)

    connection = connect_readonly(args.database)
    count = 0
    try:
        for record in iter_rows(connection, args.table, args):
            heading = (
                str(record.get("scientific_name"))
                or str(record.get("canonical_name"))
                or str(record.get("speciedex_id"))
                or f"Record {count + 1}"
            )
            document.add_heading(heading, level=1)
            table = document.add_table(rows=0, cols=2)
            table.style = "Table Grid"
            for name, value in preferred_fields(record):
                cells = table.add_row().cells
                cells[0].text = str(name)
                cells[1].text = str(value)
            count += 1
            if args.page_break_every and count % args.page_break_every == 0:
                document.add_page_break()
    finally:
        connection.close()

    document.save(args.output)

    print(json.dumps({
        "output": args.output.as_posix(),
        "records_exported": count,
        "format": "docx",
    }, indent=2, sort_keys=True))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
